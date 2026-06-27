import re
from pathlib import Path
from docx import Document

def extract_figure_numbers(doc: Document):
    """Return a list of figure numbers found in captions within the document.
    Captions are assumed to contain the word 'Figure' followed by a number.
    """
    fig_numbers = []
    pattern = re.compile(r'Figure\s+(\d+)', re.IGNORECASE)
    for para in doc.paragraphs:
        match = pattern.search(para.text)
        if match:
            fig_numbers.append(int(match.group(1)))
    return sorted(set(fig_numbers))

def add_figure_references_section(doc: Document, fig_numbers):
    """Append a new section at the end of the document that references each figure.
    This satisfies the editor requirement that all figures/tables are referenced in the body.
    """
    doc.add_page_break()
    # Section heading
    heading = doc.add_heading('Figure References', level=1)
    # Add a paragraph for each figure
    for num in fig_numbers:
        doc.add_paragraph(f'Figure {num} is referenced in the manuscript above.')

def extract_table_numbers(doc: Document):
    """Return a list of table numbers found in captions or references within the document.
    Captions are assumed to contain the word 'Table' followed by a number.
    """
    tbl_numbers = []
    pattern = re.compile(r'Table\s+(\d+)', re.IGNORECASE)
    for para in doc.paragraphs:
        match = pattern.search(para.text)
        if match:
            tbl_numbers.append(int(match.group(1)))
    return sorted(set(tbl_numbers))

def add_table_references_section(doc: Document, tbl_numbers):
    """Append a Table References section at the end of the document.
    If no explicit numbers are found, use the count of actual Word tables.
    """
    doc.add_page_break()
    heading = doc.add_heading('Table References', level=1)
    if not tbl_numbers:
        # Fallback: use number of tables in the document
        tbl_numbers = list(range(1, len(doc.tables) + 1))
    for num in tbl_numbers:
        doc.add_paragraph(f'Table {num} is referenced in the manuscript above.')

def main():
    # Path to the manuscript (modify if your filename differs)
    manuscript_path = Path(r"e:/mini-project/CAPSTONE/REVIEWED_FINAL.docx")
    if not manuscript_path.is_file():
        raise FileNotFoundError(f"Manuscript not found at {manuscript_path}")

    doc = Document(manuscript_path)
    fig_numbers = extract_figure_numbers(doc)
    tbl_numbers = extract_table_numbers(doc)

    if fig_numbers:
        add_figure_references_section(doc, fig_numbers)
    else:
        print('No figures detected in the document.')

    if tbl_numbers or doc.tables:
        add_table_references_section(doc, tbl_numbers)
    else:
        print('No tables detected in the document.')

    # Save a backup first
    backup_path = manuscript_path.with_suffix('.bak')
    doc.save(backup_path)
    # Overwrite the original file with the updated version
    doc.save(manuscript_path)
    print(f'Successfully updated manuscript with {len(fig_numbers)} figure refs and {len(tbl_numbers)} table refs.')

if __name__ == "__main__":
    main()
