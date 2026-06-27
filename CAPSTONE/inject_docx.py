from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'e:\mini-project\CAPSTONE\submitted.docx')

def insert_paragraph_after(ref_para, text, style_name):
    """Insert a new paragraph immediately after ref_para."""
    new_p_el = OxmlElement('w:p')
    ref_para._p.addnext(new_p_el)
    # Find the newly inserted paragraph object
    for p in doc.paragraphs:
        if p._p is new_p_el:
            p.style = doc.styles[style_name]
            p.text = text
            return p
    return None

# Locate the last Conclusion body paragraph (index 232) and References paragraph (index 235)
paras = doc.paragraphs
conclusion_last = paras[232]
print(f"Inserting after para[232]: {conclusion_last.text[:80]!r}")

data_availability_body = (
    "This study is a systematic literature review. All data analyzed in this paper "
    "are drawn from the twenty publicly available research papers cited in the "
    "References section. No novel datasets were generated or analyzed beyond those "
    "reported in the cited works. Readers are referred to the original papers for "
    "the underlying experimental data and results."
)
funding_body = "This research received no external funding."

# Insert in REVERSE order — each insertion goes directly after conclusion_last,
# so the final order will be: Data Availability heading → body → Funding heading → body
insert_paragraph_after(conclusion_last, funding_body,                  'Body Text')
insert_paragraph_after(conclusion_last, 'Funding',                     'Heading 1')
insert_paragraph_after(conclusion_last, data_availability_body,        'Body Text')
insert_paragraph_after(conclusion_last, 'Data Availability Statement', 'Heading 1')

doc.save(r'e:\mini-project\CAPSTONE\submitted.docx')
print("Done — submitted.docx updated successfully.")
